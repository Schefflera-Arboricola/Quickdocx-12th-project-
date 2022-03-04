from docx import Document
import os
import csv
import pickle

def add_sample():          #function to make new samples 
    n=input('Enter docfile name (without extention) : ') 
    loc=input('Enter location of the docfile : ')
    k=int(input('How many replacing items are there ? '))   # getting numer of variables from user
    d={}
    for i in range(k):
        m=input('Enter item name: ')           # getting variable names from user
        d[i]=m
    wo.writerow((n,loc,d))         # adding sample's name,location and dictionary of the format :
    print('Added',n,'as a Sample !! ')  #{1:<variable 1>,2:<variable  2>,3:<variable 3>.....}

def del_sample():        #function to delete a samples 
    sname=input('Enter docfile name (without extention) : ')
    i1=[]
    for i in ro:       #extracting records of all docfiles from samples.csv except the with name as <sname> 
        if i[0]!=sname:
            i1+=[i]
    f1=open('samples1.csv','w')    #creating a copy of samples.csv but without record of <sname> file
    w=csv.writer(f1)
    w.writerows(i1)
    f1.close()
    f.close()
    print('Sample',sname,' deleted')
    os.remove('samples.csv')          # deleting samples.csv 
    os.rename('samples1.csv','samples.csv')    #renaming old file
    
def new_doc(sample,new_doc,dic) :     
    doc = Document(sample)
    a=doc.paragraphs
    p1=[]
    for i in a:    
        a1=i.text.split('{')
        for j in a1:
            if j.isdigit()==True:
                index=a1.index(j)
                a1[index]=dic[int(j)]
        a1=' '.join(a1)
        p1+=[a1] 
    doc= Document()
    for i in p1:
        doc.add_paragraph(i)
    doc.save(new_doc)    
    print(new_doc,'Created.')
    
def use_sample():  
    sname=input('Enter sample docfile name (without extention) : ')
    s=filenames()
    i1=[]
    for i in ro:
        i1+=[i[2]]
    del i1[0]               #everytime open and close the file ?!?
    
    while sname in s and sname !='b':
        sname+='.docx'
        print('Enter 0 to stop making documents of this ')
        name=input('Enter the new docfile\'s name (without extention) you want to create : ')
        while name!='0':
            name+='.docx'
            d={}
            for num in i1:
                val=input(i1[num])            
                d[num]=val
            new_doc(sname,name,d)
            name=input('\nEnter the new docfile\'s name (without extention) you want to create : ')
        print('Enter sample name to use another sample\nEnter \'b\' to end.')
        sname=input('\n\nEnter sample docfile name (without extention) : ')
    else:
        if sname not in s:
            print('Sorry :( File not found. ')

def info():
    with open('documentation.dat','rb') as s:
        a=pickle.load(s)    
        print(a)
        
def show_samples():   
    i1=filenames()
    for i in i1:
        print(i)
        
def filenames():
    i1=[]
    for i in ro:
        i1+=[i[0]]
    del i1[0]
    return i1

#program
f=open('samples.csv','r+',newline='')
wo=csv.writer(f)
ro=csv.reader(f)
print('Enter 0 to end the program\nEnter 1 to create document(s)\nEnter 2 to add sample\nEnter 3 to see all samples\nEnter 4 to delete a sample\nEnter 5 for help')
c=int(input('0/1/2/3/4/5 : '))
while c!=0:
    if c==1:
        use_sample()   
    elif c==2:            
        add_sample()
    elif c==3:           
        show_samples()
    elif c==4:             
        del_sample()
    else:
        info()
    c=int(input('0/1/2/3/4/5 : '))
f.close()
